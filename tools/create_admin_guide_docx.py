from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "GTournament1_ADMIN_USER_GUIDE.docx"
NAVY = RGBColor(20, 32, 55)
RED = RGBColor(190, 35, 31)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(86, 96, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FCE8E7"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # Keep table rows accessible to screen readers and stable across page breaks.
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_run(run, size=11, color=None, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_text(doc, text, size=11, color=None, bold=False, italic=False, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size, color, bold, italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r, 11, NAVY)
    return p


def add_step(doc, number, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"{number}. {title}: ")
    set_run(r, 11, RED, True)
    r = p.add_run(detail)
    set_run(r, 11, NAVY)


def add_callout(doc, label, text, fill=LIGHT_RED):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run(r, 10.5, RED, True)
    r = p.add_run(text)
    set_run(r, 10.5, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, {1: 16, 2: 13, 3: 12}[level], BLUE if level < 3 else NAVY, True)
    return p


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("GTournament1 | คู่มือ Admin และทีมงาน")
    set_run(r, 9, GRAY)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)
add_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Cover / masthead.
add_text(doc, "GTournament1", 12, RED, True, after=8)
add_text(doc, "คู่มือการใช้งาน Admin และทีมงาน", 27, NAVY, True, after=4)
add_text(doc, "คู่มือสำหรับตรวจใบสมัคร จัดการ Season และดูแลข้อพิพาทการแข่งขัน", 13, GRAY, after=18)
meta = doc.add_table(rows=3, cols=2)
set_table_geometry(meta, [1800, 7560])
meta_data = [("ระบบ", "GTournament1 Production"), ("กลุ่มผู้ใช้", "Admin และ Staff"), ("ฉบับเอกสาร", "1.0 | 18 สิงหาคม 2026")]
for row, (a, b) in zip(meta.rows, meta_data):
    set_cell_shading(row.cells[0], LIGHT_BLUE)
    set_cell_shading(row.cells[1], "F7F9FB")
    row.cells[0].text = ""
    row.cells[1].text = ""
    p = row.cells[0].paragraphs[0]
    set_run(p.add_run(a), 10.5, NAVY, True)
    p = row.cells[1].paragraphs[0]
    set_run(p.add_run(b), 10.5, NAVY)
add_callout(doc, "สรุปสถานะ:", "ตรวจสอบเส้นทาง Admin, Staff และ Match Queue แล้ว ระบบตรวจ Role ที่ฝั่งเซิร์ฟเวอร์ และเฉพาะ Admin เท่านั้นที่จัดการ Season ได้", "EAF4EA")

heading(doc, "1. เริ่มต้นใช้งาน", 1)
add_step(doc, 1, "เปิดเว็บไซต์", "เข้าเว็บไซต์ GTournament1 แล้วกดเมนู เฉพาะทีมงาน ที่มุมขวาบน")
add_step(doc, 2, "เข้าสู่ระบบ", "กรอกอีเมลและรหัสผ่านทีมงาน จากนั้นกด เข้าสู่ระบบทีมงาน")
add_step(doc, 3, "ตรวจ Role", "Admin จะไปหน้า Admin ส่วน Staff จะไปหน้าตรวจใบสมัคร หากไม่มี Role ระบบจะปฏิเสธการเข้าถึง")
add_callout(doc, "ความปลอดภัย:", "อย่าใส่รหัสผ่านลงในคู่มือหรือส่งต่อในแชตสาธารณะ ควรเปลี่ยนรหัสผ่านเริ่มต้นทันทีหลังการเข้าสู่ระบบครั้งแรก")

heading(doc, "2. เมนูและหน้าที่", 1)
table = doc.add_table(rows=1, cols=3)
set_table_geometry(table, [2200, 2800, 4360])
for cell, text in zip(table.rows[0].cells, ["เมนู", "เส้นทาง", "ทำอะไรได้บ้าง"]):
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    set_run(cell.paragraphs[0].add_run(text), 10.5, NAVY, True)
rows = [
    ("เฉพาะทีมงาน", "?view=auth / admin / staff", "เข้าสู่ระบบทีมงาน และไปยังหน้าตาม Role"),
    ("Admin", "?view=admin", "สร้างและแก้ไขข้อมูล Season"),
    ("ตรวจใบสมัคร", "?view=staff", "ค้นหา เปิดสลิป อนุมัติ หรือปฏิเสธใบสมัคร"),
    ("Match Queue", "?view=staff-matches", "ตรวจหลักฐานและตัดสินข้อพิพาทการแข่งขัน"),
    ("Health check", "?view=health", "ตรวจสถานะ environment, session และ Supabase"),
]
for vals in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, vals):
        cell.text = ""
        set_run(cell.paragraphs[0].add_run(text), 10, NAVY)

heading(doc, "3. Admin: สร้างหรือแก้ไข Season", 1)
add_text(doc, "ไปที่เมนู เฉพาะทีมงาน แล้วเลือกหน้า Admin / จัดการ Season และการรับสมัคร", 11, NAVY)
heading(doc, "ช่องข้อมูลที่ต้องกรอก", 2)
fields = [
    ("Season ID", "เว้นว่างเมื่อสร้างใหม่ ใส่ ID เดิมเมื่อต้องการแก้ไข Season ที่มีอยู่"),
    ("ชื่อ Season", "เช่น GTournament1 Season 01"),
    ("คำอธิบาย", "รายละเอียดสั้น ๆ ของรายการแข่งขัน"),
    ("สถานะ", "ร่าง, เปิดรับสมัคร, ปิดรับสมัคร, กำลังแข่งขัน หรือจบแล้ว"),
    ("Capacity", "จำนวนผู้สมัครสูงสุด กรอกเอง เช่น 32"),
    ("ค่าสมัคร", "จำนวนเงินค่าสมัครต่อคน กรอกเอง เช่น 25"),
    ("ช่วงเวลา", "วันและเวลาเปิด-ปิดรับสมัคร"),
    ("PromptPay", "ชื่อบัญชี หมายเลข และยอดที่ต้องชำระ"),
]
field_table = doc.add_table(rows=1, cols=2)
set_table_geometry(field_table, [2100, 7260])
for cell, text in zip(field_table.rows[0].cells, ["ช่อง", "วิธีใช้"]):
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    set_run(cell.paragraphs[0].add_run(text), 10.5, NAVY, True)
for a, b in fields:
    cells = field_table.add_row().cells
    cells[0].text = ""
    cells[1].text = ""
    set_run(cells[0].paragraphs[0].add_run(a), 10, NAVY, True)
    set_run(cells[1].paragraphs[0].add_run(b), 10, NAVY)
add_step(doc, 1, "กรอกข้อมูล", "กรอกชื่อ Season, Capacity, ค่าสมัคร และข้อมูลการชำระเงิน")
add_step(doc, 2, "เลือกสถานะ", "ใช้ ร่าง ระหว่างเตรียมข้อมูล และเลือก เปิดรับสมัครเมื่อพร้อม")
add_step(doc, 3, "บันทึก", "กด บันทึก Season แล้วรอข้อความยืนยัน")
add_callout(doc, "ข้อจำกัดปัจจุบัน:", "หน้า Admin ยังไม่มีรายการ Season พร้อมปุ่มแก้ไขอัตโนมัติ หากแก้รายการเดิมต้องใช้ Season ID ของรายการนั้น")

heading(doc, "4. Staff: ตรวจใบสมัครและสลิป", 1)
add_step(doc, 1, "เปิดคิว", "เลือก เฉพาะทีมงาน แล้วไปหน้าตรวจใบสมัคร")
add_step(doc, 2, "กรองรายการ", "ใช้ รอตรวจ, อนุมัติแล้ว, ไม่ผ่าน หรือ ทั้งหมด")
add_step(doc, 3, "ค้นหา", "พิมพ์ชื่อผู้สมัครหรือชื่อที่ใช้แข่งขัน แล้วกดค้นหา")
add_step(doc, 4, "ตรวจหลักฐาน", "กด เปิดสลิป เพื่อตรวจหลักฐานใน private storage")
add_step(doc, 5, "ตัดสิน", "กด อนุมัติ หรือ ปฏิเสธ โดยการปฏิเสธต้องใส่เหตุผล")
add_text(doc, "ทุกการอนุมัติและปฏิเสธจะถูกบันทึกไว้ใน Audit Log", 11, NAVY, True)

heading(doc, "5. Match Queue: ตรวจข้อพิพาท", 1)
add_step(doc, 1, "เปิด Match Queue", "กด Match Queue บนเมนูด้านบน")
add_step(doc, 2, "ตรวจข้อมูล", "ดูรอบแข่งขัน สกอร์ ผู้ส่งผล และสถานะข้อพิพาท")
add_step(doc, 3, "ตรวจหลักฐาน", "เปิด Screenshot และ Screen Recording ที่แนบมา")
add_step(doc, 4, "ตัดสินผล", "กดยืนยันผลตามหลักฐาน หรือยกเลิกผลแมตช์")
add_callout(doc, "ข้อควรระวัง:", "ตัดสินจากหลักฐานและหมายเหตุให้ครบ ระบบจะไม่อนุญาตให้ตัดสินแมตช์เดิมซ้ำ")

heading(doc, "6. สิทธิ์การใช้งาน", 1)
role_table = doc.add_table(rows=1, cols=4)
set_table_geometry(role_table, [3300, 2020, 2020, 2020])
for cell, text in zip(role_table.rows[0].cells, ["ความสามารถ", "ผู้เล่น", "Staff", "Admin"]):
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    set_run(cell.paragraphs[0].add_run(text), 10, NAVY, True)
permissions = [
    ("ดูหน้าเว็บไซต์ทั่วไป", "ได้", "ได้", "ได้"),
    ("ตรวจใบสมัครและสลิป", "ไม่ได้", "ได้", "ได้"),
    ("ตรวจ Match Queue", "ไม่ได้", "ได้", "ได้"),
    ("สร้าง/แก้ไข Season", "ไม่ได้", "ไม่ได้", "ได้"),
    ("เข้าหน้า Admin", "ไม่ได้", "ไม่ได้", "ได้"),
]
for vals in permissions:
    cells = role_table.add_row().cells
    for cell, text in zip(cells, vals):
        cell.text = ""
        set_run(cell.paragraphs[0].add_run(text), 10, NAVY, text == "ได้" and vals[0] in ("สร้าง/แก้ไข Season", "เข้าหน้า Admin"))

heading(doc, "7. ตรวจสอบเมื่อพบปัญหา", 1)
checks = [
    "ข้อความ 403 Forbidden: บัญชีไม่มี Role ที่เพียงพอ",
    "ระบบขัดข้องชั่วคราว: เปิด ?view=health ตรวจ environment, session และ Supabase",
    "บันทึก Season ไม่ได้: ตรวจว่าเป็น Role admin และกรอกชื่อ Season แล้ว",
    "ไม่มีรายการใบสมัครหรือข้อพิพาท: ระบบทำงานปกติ แต่อาจยังไม่มีข้อมูลตามตัวกรอง",
    "เมนูไม่เปลี่ยน: กด Ctrl + F5 เพื่อล้าง cache หน้าเว็บ",
]
for item in checks:
    add_bullet(doc, item)

heading(doc, "8. Checklist ก่อนเปิด Season", 1)
for item in [
    "ตรวจชื่อ Season และคำอธิบาย",
    "ตรวจ Capacity และค่าสมัคร",
    "ตรวจช่วงเวลาเปิด-ปิดรับสมัคร",
    "ตรวจ PromptPay และยอดที่ต้องชำระ",
    "ตั้งสถานะเป็น เปิดรับสมัคร เมื่อพร้อมเท่านั้น",
    "ทดสอบส่งข้อมูลผู้สมัครหนึ่งรายการ",
    "ตรวจว่า Staff เปิดสลิปและเห็นคิวตรวจได้",
]:
    add_bullet(doc, item)

add_callout(doc, "สถานะการตรวจล่าสุด:", "เว็บจริงตอบสนองปกติ, Health check ผ่าน, Role guard ของ Admin/Staff ทำงาน และ RLS สำหรับ Admin จัดการ Season ถูกติดตั้งแล้ว", "EAF4EA")

doc.core_properties.title = "GTournament1 คู่มือการใช้งาน Admin และทีมงาน"
doc.core_properties.subject = "คู่มือจัดการ Season ใบสมัคร และข้อพิพาท"
doc.core_properties.author = "GTournament1"
doc.core_properties.keywords = "GTournament1, Admin, Staff, Season, Match Queue"
doc.save(OUT)
print(OUT)
